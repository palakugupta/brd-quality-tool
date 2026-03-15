from typing import Optional, Dict

from fastapi import APIRouter, UploadFile, File, HTTPException, status
from pydantic import BaseModel
from ..database import get_connection
from pypdf import PdfReader
import docx

from ..preprocessing.requirement_blocks import extract_requirement_blocks

from ..models import (
    insert_document,
    create_brd_chunks,
    get_document,
    get_chunks_for_brd,
    get_findings_for_brd,
    create_analysis_run,
    finalize_analysis_run,
)

from ..detectors import (
    different_data,
    incomplete_data,
    hallucination,
    depth_mismatch,
    duplicate_data,
    platform_constraints,
    process_flow_validator,
    terminology_drift,
    missing_process_steps,
    business_rule_violation,
    role_responsibility_violation,
    organization_mismatch,
    process_dependency_validator,
)

router = APIRouter(prefix="/api", tags=["analysis"])


# ─────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────

def extract_text_from_pdf(file: UploadFile) -> str:
    try:
        reader = PdfReader(file.file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to read PDF: {e}",
        )


def extract_text_from_docx(file: UploadFile) -> str:
    try:
        document = docx.Document(file.file)
        lines = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                lines.append(
                    " | ".join(c.text.replace("\n", " ").strip() for c in row.cells)
                )
        return "\n".join(lines)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to read DOCX: {e}",
        )


def extract_text(file: UploadFile) -> str:
    filename = (file.filename or "").lower()
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file)
    if filename.endswith(".docx"):
        return extract_text_from_docx(file)
    if filename.endswith(".txt"):
        file.file.seek(0)
        return file.file.read().decode("utf-8", errors="ignore")
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Unsupported file type. Upload PDF, DOCX, or TXT.",
    )


# ─────────────────────────────────────────────
# DOCUMENT UPLOAD
# ─────────────────────────────────────────────

@router.post("/upload-documents")
async def upload_documents(
    input_sow: Optional[UploadFile] = File(None),
    input_mom: Optional[UploadFile] = File(None),
    output_brd: UploadFile = File(...),
):
    if output_brd.size == 0:
        raise HTTPException(400, "Output BRD file is empty")

    sow_doc_id = None
    mom_doc_id = None
    sow_line_count = None
    mom_line_count = None

    if input_sow:
        sow_text = extract_text(input_sow)
        sow_doc_id, sow_line_count = insert_document(
            "input_sow", input_sow.filename, sow_text
        )

    if input_mom:
        mom_text = extract_text(input_mom)
        mom_doc_id, mom_line_count = insert_document(
            "input_mom", input_mom.filename, mom_text
        )

    brd_text = extract_text(output_brd)
    if not brd_text.strip():
        raise HTTPException(400, "BRD file is empty")

    brd_doc_id, brd_line_count = insert_document(
        "output_brd", output_brd.filename, brd_text
    )
    chunks_created = create_brd_chunks(brd_doc_id, brd_text, chunk_size=120)

    print("UPLOAD COMPLETE")
    print("SOW lines:", sow_line_count)
    print("MoM lines:", mom_line_count)
    print("BRD lines:", brd_line_count)
    print("Chunks:", chunks_created)

    return {
        "message": "Documents uploaded successfully",
        "sow": {"doc_id": sow_doc_id, "line_count": sow_line_count},
        "mom": {"doc_id": mom_doc_id, "line_count": mom_line_count},
        "brd": {
            "doc_id": brd_doc_id,
            "line_count": brd_line_count,
            "chunks_created": chunks_created,
        },
    }


# ─────────────────────────────────────────────
# ANALYSIS REQUEST
# ─────────────────────────────────────────────

class AnalysisRequest(BaseModel):
    sow_doc_id: int
    mom_doc_id: Optional[int] = None
    brd_doc_id: int


# ─────────────────────────────────────────────
# LOAD DOCUMENT TEXTS
# ─────────────────────────────────────────────

def _load_texts(payload: AnalysisRequest):
    sow_doc = get_document(payload.sow_doc_id)
    brd_doc = get_document(payload.brd_doc_id)

    if not sow_doc or not brd_doc:
        raise HTTPException(400, "Invalid document id")

    sow_text = sow_doc["full_text"] or ""
    brd_text = brd_doc["full_text"] or ""
    mom_text = ""

    if payload.mom_doc_id:
        mom_doc = get_document(payload.mom_doc_id)
        if mom_doc:
            mom_text = mom_doc["full_text"] or ""

    chunks = get_chunks_for_brd(payload.brd_doc_id)
    blocks = extract_requirement_blocks(brd_text)

    print("TEXT LOAD CHECK")
    print("SOW length:", len(sow_text))
    print("MoM length:", len(mom_text))
    print("BRD length:", len(brd_text))
    print("Chunks:", len(chunks))
    print("Blocks extracted:", len(blocks))

    return sow_text, mom_text, brd_text, chunks, blocks


# ─────────────────────────────────────────────
# BUILD SUMMARY
# ─────────────────────────────────────────────

def _build_summary(brd_doc_id: int):
    findings = get_findings_for_brd(brd_doc_id)

    summary: Dict[str, int] = {}
    for f in findings:
        summary.setdefault(f["error_type"], 0)
        summary[f["error_type"]] += 1

    total_findings = len(findings)
    clean_pct = 0
    conn = None

    try:
        conn = get_connection()
        cur = conn.cursor()

        row = cur.execute(
            "SELECT COALESCE(line_count, 0) FROM documents WHERE doc_type='output_brd' AND doc_id=?",
            (brd_doc_id,),
        ).fetchone()
        total_lines = row[0] if row else 0

        flagged = cur.execute(
            """
            SELECT COUNT(DISTINCT f.line_number)
            FROM findings f
            JOIN chunks c ON f.chunk_id = c.chunk_id
            WHERE f.line_number IS NOT NULL
              AND c.doc_id = ?
            """,
            (brd_doc_id,),
        ).fetchone()[0]

        clean_pct = ((total_lines - flagged) / total_lines * 100) if total_lines else 0

    except Exception as e:
        print("Summary error:", e)
        clean_pct = 0

    finally:
        if conn:
            conn.close()

    return {
        "summary": summary,
        "total_findings": total_findings,
        "coverage_score": round(clean_pct, 2),
        "findings": findings,
    }


# ─────────────────────────────────────────────
# FULL ANALYSIS
# ─────────────────────────────────────────────

@router.post("/run-full-analysis")
async def run_full_analysis(payload: AnalysisRequest):

    sow_text, mom_text, brd_text, chunks, blocks = _load_texts(payload)

    # ── Clear previous findings for this BRD ──
    conn = get_connection()
    conn.execute(
        "DELETE FROM findings WHERE chunk_id IN (SELECT chunk_id FROM chunks WHERE doc_id = ?)",
        (payload.brd_doc_id,),
    )
    conn.commit()
    conn.close()

    run_id = create_analysis_run(
        sow_doc_id=payload.sow_doc_id,
        mom_doc_id=payload.mom_doc_id,
    )

    print("Running document-level detectors...")

    # ── Run ONCE on full BRD ──────────────────────────────────
    try:
        different_data.detect(sow_text, mom_text, brd_text, chunks)
    except Exception as e:
        print("different_data error:", e)

    try:
        incomplete_data.detect(sow_text, brd_text, chunks)
    except Exception as e:
        print("incomplete_data error:", e)

    try:
        hallucination.detect(sow_text, mom_text, brd_text, chunks)
    except Exception as e:
        print("hallucination error:", e)

    try:
        depth_mismatch.detect(sow_text, brd_text, chunks)
    except Exception as e:
        print("depth_mismatch error:", e)

    try:
        duplicate_data.detect(brd_text, chunks)
    except Exception as e:
        print("duplicate_data error:", e)

    try:
        terminology_drift.detect(sow_text, mom_text, brd_text, chunks)
    except Exception as e:
        print("terminology_drift error:", e)

    try:
        missing_process_steps.detect(sow_text, brd_text, chunks)
    except Exception as e:
        print("missing_process_steps error:", e)

    try:
        organization_mismatch.detect(sow_text, mom_text, brd_text, chunks)
    except Exception as e:
        print("organization_mismatch error:", e)

    try:
        process_dependency_validator.detect(brd_text, chunks)
    except Exception as e:
        print("process_dependency_validator error:", e)

    try:
        process_flow_validator.detect(brd_text, chunks)
    except Exception as e:
        print("process_flow_validator error:", e)

    print("Running block-level detectors...")

    if not blocks:
        blocks = [{"text": brd_text}]

    # ── Run PER BLOCK ─────────────────────────────────────────
    # platform_constraints and role_responsibility are pattern matchers
    # that benefit from granular block text.
    # business_rule_violation now needs sow_text + mom_text so runs once on full BRD.
    try:
        business_rule_violation.detect(
            brd_text,
            chunks,
            sow_text=sow_text,
            mom_text=mom_text,
        )
    except Exception as e:
        print("business_rule_violation error:", e)

    for block in blocks:
        block_text = block.get("text", "")
        if not block_text.strip():
            continue

        try:
            platform_constraints.detect(block_text, chunks)
        except Exception as e:
            print("platform_constraints error:", e)

        try:
            role_responsibility_violation.detect(block_text, chunks)
        except Exception as e:
            print("role_responsibility_violation error:", e)

    summary = _build_summary(payload.brd_doc_id)

    finalize_analysis_run(
        run_id=run_id,
        total_findings=summary["total_findings"],
        coverage_score=summary["coverage_score"],
    )

    return summary